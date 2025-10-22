from docx import Document
import sys

docx_path = r'knowledge\CHRIS_NATTRESS_RESUME.docx'
output_path = 'resume_extracted.txt'

try:
    doc = Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Also extract text from tables if any
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    
    content = '\n'.join(full_text)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    print(f"Successfully extracted text to {output_path}")
    print(f"Total characters: {len(content)}")
except Exception as e:
    print(f"Error: {e}", file=sys.stderr)
    sys.exit(1)
